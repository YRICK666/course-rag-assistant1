from __future__ import annotations

import hashlib
import io
import json
import logging
import re
from urllib.parse import quote
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from pathlib import Path, PureWindowsPath
from typing import Literal

from docx import Document
from docx.shared import Pt
from fastapi import FastAPI, File, HTTPException, Response, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field


logger = logging.getLogger(__name__)


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SRC_DIR = PROJECT_ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from ai_settings import load_ai_settings, public_ai_settings, save_ai_settings  # noqa: E402
from llm_deepseek import DeepSeekError, extract_snippet_keywords, is_deepseek_configured  # noqa: E402
from material_manager import list_materials, material_stats  # noqa: E402
from ppt_converter import find_libreoffice_executable  # noqa: E402
from retriever import fetch_all_records, get_collection, list_indexed_sources, reset_chroma_client_cache  # noqa: E402
from services.index_service import (  # noqa: E402
    build_selected_materials_index,
    rebuild_all_materials_index,
)
from services.learning_service import (  # noqa: E402
    generate_all_materials_overview,
    generate_current_material_overview,
    generate_study_guide,
)
from services.material_service import (  # noqa: E402
    convert_subject_ppt_material,
    rename_subject_material,
    save_uploaded_materials,
    soft_delete_subject_material,
)
from services.qa_service import ask_course_question  # noqa: E402
from services.scope_service import indexed_source_path_set  # noqa: E402
from .services.longform_service import generate_longform_analysis  # noqa: E402
from .services.self_test_service import generate_self_test  # noqa: E402
from subject_store import SUBJECTS_DIR, create_subject, get_subject_paths, list_subjects  # noqa: E402
from .db import init_db  # noqa: E402
from .qa_history import (  # noqa: E402
    create_qa_record,
    delete_qa_record,
    ensure_qa_history_schema,
    get_qa_record,
    list_qa_records,
    qa_history_to_dict,
)


app = FastAPI(title="Course Materials RAG API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://127.0.0.1:5173",
        "http://localhost:5173",
        "http://127.0.0.1:5174",
        "http://localhost:5174",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_PREVIEW_CONVERSION_LOCKS: dict[str, threading.Lock] = {}
_PREVIEW_CONVERSION_LOCKS_GUARD = threading.Lock()
AI_DISABLED_WARNING = "AI 已关闭，仍可查看资料来源和页面预览。"


@app.on_event("startup")
def on_startup():
    init_db()
    ensure_qa_history_schema()


class QaRequest(BaseModel):
    question: str
    source_filters: list[str] = Field(default_factory=list)
    top_k: int = 8
    use_deepseek: bool = True


class SelfTestTypeConfig(BaseModel):
    type: Literal["choice", "fill", "essay"]
    count: int = Field(default=0, ge=0, le=20)


class SelfTestRequest(BaseModel):
    source_filters: list[str] = Field(default_factory=list)
    type_configs: list[SelfTestTypeConfig] = Field(default_factory=list)
    answer_mode: Literal["inline", "end", "dual"] = "inline"


class LongformRequest(BaseModel):
    source_filters: list[str] = Field(default_factory=list)
    longform_type: Literal["analysis", "study_notes", "report", "review", "outline"] = "analysis"
    target_length: int = Field(default=3000, ge=500, le=15000)
    include_sources: bool = True
    strategy: Literal["staged"] = "staged"
    user_instruction: str = Field(default="", max_length=1000)


class SnippetKeywordsRequest(BaseModel):
    text: str = Field(default="", max_length=6000)


class AiSettingsRequest(BaseModel):
    enabled: bool
    provider: Literal["deepseek", "openai_compatible"] = "deepseek"
    base_url: str = ""
    model: str = ""
    api_key: str | None = None


class IndexRequest(BaseModel):
    mode: Literal["reset", "update"] = "update"
    scope: Literal["all", "selected"] = "selected"
    files: list[str] = Field(default_factory=list)


class StudyGuideRequest(BaseModel):
    source_filters: list[str] = Field(default_factory=list)
    use_deepseek: bool = True
    force_refresh: bool = False


class OverviewRequest(BaseModel):
    source_filters: list[str] = Field(default_factory=list)
    use_deepseek: bool = True
    force_refresh: bool = False


class CreateSubjectRequest(BaseModel):
    name: str


class DeleteMaterialRequest(BaseModel):
    relative_path: str


class RenameMaterialRequest(BaseModel):
    old_relative_path: str
    new_relative_path: str


class ConvertPptRequest(BaseModel):
    relative_path: str


class ExportSelfTestRequest(BaseModel):
    subject: str
    scope_label: str = ""
    generated_at: str = ""
    content: str = ""
    sources: list[dict] = Field(default_factory=list)
    include_sources: bool = False
    filename: str = ""


class DocxExportRequest(BaseModel):
    title: str
    subject: str
    scope_label: str = ""
    generated_at: str = ""
    content: str = ""
    sources: list[dict] = Field(default_factory=list)
    include_sources: bool = True
    filename_prefix: str = "export"
    filename: str = ""


class ApiUploadedFile:
    def __init__(self, name: str, data: bytes) -> None:
        self.name = name
        self._data = data

    def getbuffer(self) -> bytes:
        return self._data


def resolve_existing_subject_paths(subject: str):
    subjects = list_subjects()
    if subject not in subjects:
        raise HTTPException(status_code=404, detail="Subject not found.")
    return get_subject_paths(subject)


def validate_new_subject_name(name: str) -> str:
    normalized = " ".join(name.strip().split())
    if not normalized:
        raise HTTPException(status_code=400, detail="科目名称不能为空。")
    if any(separator in normalized for separator in ("/", "\\")):
        raise HTTPException(status_code=400, detail="科目名称不能包含路径分隔符。")
    if normalized in {".", ".."} or ".." in normalized:
        raise HTTPException(status_code=400, detail="科目名称不能包含 '..'。")

    subject_path = Path(normalized)
    windows_subject_path = PureWindowsPath(normalized)
    if subject_path.is_absolute() or windows_subject_path.is_absolute() or windows_subject_path.drive:
        raise HTTPException(status_code=400, detail="科目名称不能是绝对路径。")

    subjects_root = SUBJECTS_DIR.resolve()
    target_root = (subjects_root / normalized).resolve()
    if target_root != subjects_root and subjects_root not in target_root.parents:
        raise HTTPException(status_code=400, detail="科目目录必须位于 subjects 根目录内。")
    return normalized


def resolve_material_file(subject_paths, source_path: str) -> Path:
    source = Path(source_path)
    if not source_path.strip() or source.is_absolute() or ".." in source.parts:
        raise HTTPException(status_code=400, detail="Invalid source_path.")

    materials_root = subject_paths.materials_dir.resolve()
    target = (materials_root / source).resolve()
    if target != materials_root and materials_root not in target.parents:
        raise HTTPException(status_code=400, detail="Invalid source_path.")
    return target


def is_inside_directory(path: Path, root: Path) -> bool:
    resolved_path = path.resolve()
    resolved_root = root.resolve()
    return resolved_path == resolved_root or resolved_root in resolved_path.parents


def resolve_preview_material_file(subject_paths, source_path: str) -> Path:
    material_path = resolve_material_file(subject_paths, source_path)
    if material_path.exists():
        return material_path

    if material_path.suffix.lower() == ".ppt":
        materials_root = subject_paths.materials_dir.resolve()
        for candidate in sorted(material_path.parent.glob(f"{material_path.stem}*.pptx")):
            if candidate.is_file() and is_inside_directory(candidate, materials_root):
                return candidate

    return material_path


def preview_cache_dir(subject_paths) -> Path:
    cache_dir = subject_paths.outputs_dir / "page_preview_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir


def preview_cache_key(source_path: str, material_path: Path, page_number: int) -> str:
    stat = material_path.stat()
    raw_key = "|".join(
        [
            source_path,
            str(material_path.resolve()),
            str(stat.st_mtime_ns),
            str(stat.st_size),
            str(page_number),
        ]
    )
    return hashlib.sha256(raw_key.encode("utf-8")).hexdigest()


def presentation_pdf_cache_key(source_path: str, material_path: Path) -> str:
    stat = material_path.stat()
    raw_key = "|".join(
        [
            source_path,
            str(material_path.resolve()),
            str(stat.st_mtime_ns),
            str(stat.st_size),
            "presentation-pdf",
        ]
    )
    return hashlib.sha256(raw_key.encode("utf-8")).hexdigest()


def presentation_pdf_cache_path(cache_dir: Path, source_path: str, material_path: Path) -> Path:
    return cache_dir / f"{presentation_pdf_cache_key(source_path, material_path)}.pdf"


def preview_conversion_lock_key(subject_paths, source_path: str, material_path: Path) -> str:
    stat = material_path.stat()
    raw_key = "|".join(
        [
            subject_paths.name,
            source_path,
            str(material_path.resolve()),
            str(stat.st_mtime_ns),
            str(stat.st_size),
        ]
    )
    return hashlib.sha256(raw_key.encode("utf-8")).hexdigest()


def preview_conversion_lock(lock_key: str) -> threading.Lock:
    with _PREVIEW_CONVERSION_LOCKS_GUARD:
        lock = _PREVIEW_CONVERSION_LOCKS.get(lock_key)
        if lock is None:
            lock = threading.Lock()
            _PREVIEW_CONVERSION_LOCKS[lock_key] = lock
        return lock


def cleanup_old_preview_cache(cache_dir: Path, keep_paths: set[Path], max_age_days: int = 30) -> None:
    try:
        resolved_cache_dir = cache_dir.resolve()
    except OSError:
        return

    resolved_keep_paths: set[Path] = set()
    for path in keep_paths:
        try:
            resolved_path = path.resolve()
        except OSError:
            continue
        if is_inside_directory(resolved_path, resolved_cache_dir):
            resolved_keep_paths.add(resolved_path)

    cutoff = time.time() - max_age_days * 24 * 60 * 60
    try:
        cache_entries = list(cache_dir.iterdir())
    except OSError:
        return

    for entry in cache_entries:
        try:
            resolved_entry = entry.resolve()
            if not is_inside_directory(resolved_entry, resolved_cache_dir):
                continue
            if resolved_entry in resolved_keep_paths:
                continue
            if not entry.is_file() or entry.suffix.lower() not in {".png", ".pdf"}:
                continue
            stat = entry.stat()
            if max(stat.st_atime, stat.st_mtime) >= cutoff:
                continue
            entry.unlink()
        except OSError:
            continue


def convert_presentation_to_cached_pdf(subject_paths, source_path: str, material_path: Path) -> Path:
    cache_dir = preview_cache_dir(subject_paths)
    target_pdf = presentation_pdf_cache_path(cache_dir, source_path, material_path)
    if target_pdf.exists():
        return target_pdf

    lock_key = preview_conversion_lock_key(subject_paths, source_path, material_path)
    with preview_conversion_lock(lock_key):
        if target_pdf.exists():
            return target_pdf

        executable = find_libreoffice_executable()
        if not executable:
            raise HTTPException(
                status_code=500,
                detail="无法生成 PPT/PPTX 预览：未检测到 LibreOffice。",
            )

        try:
            with tempfile.TemporaryDirectory(dir=str(cache_dir)) as temp_dir:
                temp_root = Path(temp_dir)
                temp_input = temp_root / material_path.name
                temp_output_dir = temp_root / "converted"
                temp_output_dir.mkdir(parents=True, exist_ok=True)
                shutil.copy2(material_path, temp_input)

                command = [
                    executable,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(temp_output_dir),
                    str(temp_input),
                ]
                result = subprocess.run(command, capture_output=True, text=True, timeout=180)
                if result.returncode != 0:
                    raise HTTPException(
                        status_code=500,
                        detail="PPT/PPTX 转 PDF 失败，请确认文件可正常打开。",
                    )

                generated_pdf = temp_output_dir / f"{temp_input.stem}.pdf"
                if not generated_pdf.exists():
                    candidates = list(temp_output_dir.glob("*.pdf"))
                    if not candidates:
                        raise HTTPException(
                            status_code=500,
                            detail="PPT/PPTX 转 PDF 失败，未生成可预览的 PDF。",
                        )
                    generated_pdf = candidates[0]

                shutil.move(str(generated_pdf), str(target_pdf))
        except HTTPException:
            raise
        except subprocess.TimeoutExpired as exc:
            raise HTTPException(status_code=500, detail="PPT/PPTX 转 PDF 超时，请稍后重试。") from exc
        except Exception as exc:
            raise HTTPException(status_code=500, detail="PPT/PPTX 转 PDF 失败，请确认文件可正常打开。") from exc

    return target_pdf


def render_pdf_page_to_png(pdf_path: Path, page_number: int, page_label: str = "页码") -> bytes:
    try:
        import fitz

        with fitz.open(str(pdf_path)) as document:
            if page_number > document.page_count:
                raise HTTPException(status_code=400, detail=f"{page_label}超出文件范围。")
            page = document.load_page(page_number - 1)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            return pixmap.tobytes("png")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail="页面预览渲染失败，请确认文件未损坏。") from exc


def api_hit(hit: dict) -> dict:
    return {
        "rank": hit.get("rank"),
        "source": hit.get("source"),
        "text": hit.get("text"),
        "metadata": hit.get("metadata", {}),
        "similarity": hit.get("similarity"),
        "hybrid_score": hit.get("hybrid_score"),
        "keyword_score": hit.get("keyword_score"),
    }


def is_index_corruption_error(message: str) -> bool:
    normalized = message.casefold()
    return any(
        marker in normalized
        for marker in (
            "error loading hnsw index",
            "hnsw segment reader",
            "hnsw",
            "segment reader",
            "compactor",
        )
    )


def is_missing_index_error(message: str) -> bool:
    normalized = message.casefold()
    return "does not exist" in normalized or "not found" in normalized


def corrupted_index_warning() -> str:
    return "当前科目的知识库索引可能已损坏，请重建知识库。"


def is_empty_index_error(message: str) -> bool:
    return "知识库为空" in message or "还没有可用知识库" in message


def qa_error_response(message: str) -> dict:
    if is_index_corruption_error(message) or "知识库索引可能已损坏" in message:
        return {
            "answer": "",
            "warning": corrupted_index_warning(),
            "hits": [],
            "error_type": "index_corrupted",
        }
    if is_empty_index_error(message):
        return {
            "answer": "",
            "warning": "当前科目还没有可用知识库，请先建立知识库。",
            "hits": [],
            "error_type": "index_empty",
        }
    return {
        "answer": "",
        "warning": "查询失败，请检查后端日志。",
        "hits": [],
        "error_type": "query_failed",
    }


def unique_nonempty_files(files: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for file_name in files:
        normalized = str(file_name).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        result.append(normalized)
    return result


def available_material_paths(subject_paths) -> set[str]:
    return {material["relative_path"] for material in list_materials(subject_paths)}


def is_unconverted_ppt(material: dict) -> bool:
    file_type = str(material.get("file_type") or "").casefold()
    file_name = str(material.get("file_name") or "").casefold()
    if file_type != ".ppt" and not file_name.endswith(".ppt"):
        return False

    conversion_status = str(material.get("conversion_status") or "")
    return not material.get("converted_pptx") or conversion_status in {
        "待转换",
        "未转换",
        "转换失败",
    }


def material_build_status(material: dict, indexed_paths: set[str] | None, index_warning: str | None) -> str:
    if indexed_paths is not None and material.get("relative_path") in indexed_paths:
        return "已建库"
    if is_unconverted_ppt(material):
        return "待转换"
    if index_warning:
        return "索引异常"
    return "未建库"


def index_health(subject_paths) -> tuple[int, str, str | None]:
    try:
        collection = get_collection(create=False, outputs_dir=subject_paths.outputs_dir)
        indexed_count = int(collection.count())
        return indexed_count, "ready" if indexed_count > 0 else "empty", None
    except Exception as exc:
        message = str(exc)
        logger.exception("Chroma 索引健康检查异常: %s", message)
        if is_missing_index_error(message):
            return 0, "empty", None
        if is_index_corruption_error(message):
            return 0, "corrupted", corrupted_index_warning()
        return 0, "error", f"Failed to read Chroma index: {message}"


def indexed_material_paths(subject_paths) -> tuple[set[str] | None, str | None]:
    try:
        collection = get_collection(create=False, outputs_dir=subject_paths.outputs_dir)
        records = fetch_all_records(collection)
    except Exception as exc:
        message = str(exc)
        logger.exception("Chroma 资料路径查询异常: %s", message)
        if is_missing_index_error(message):
            return set(), None
        if is_index_corruption_error(message):
            return None, corrupted_index_warning()
        return None, f"Failed to read Chroma index: {message}"

    indexed_paths = {
        str(record.get("metadata", {}).get("source_path"))
        for record in records
        if record.get("metadata", {}).get("source_path")
    }
    return indexed_paths, None


def index_error_response(warning: str, *, error: str | None = None, error_type: str | None = None) -> dict:
    response = {
        "success": False,
        "warning": warning,
        "error": error,
        "file_count": 0,
        "chunk_count": 0,
        "chroma_count": 0,
        "ppt_conversion": {
            "success_count": 0,
            "failure_count": 0,
            "failures": [],
        },
        "messages": [warning],
    }
    if error_type:
        response["error_type"] = error_type
    return response


def index_success_response(result: dict, request: IndexRequest) -> dict:
    ppt_conversion = result.get("ppt_conversion") or {}
    return {
        "success": bool(result.get("success")),
        "mode": request.mode,
        "scope": request.scope,
        "file_count": result.get("file_count", 0),
        "chunk_count": result.get("chunk_count", 0),
        "chroma_count": result.get("chroma_count", 0),
        "ppt_conversion": {
            "success_count": ppt_conversion.get("success_count", 0),
            "failure_count": ppt_conversion.get("failure_count", 0),
            "failures": ppt_conversion.get("failures", []),
        },
        "messages": result.get("messages") or [result.get("message", "建库完成。")],
        "warning": result.get("warning"),
    }


@app.get("/api/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/api/ai/settings")
def get_ai_settings() -> dict:
    return public_ai_settings(load_ai_settings())


@app.post("/api/ai/settings")
def update_ai_settings(request: AiSettingsRequest) -> dict:
    settings = save_ai_settings(request.dict())
    return public_ai_settings(settings)


@app.get("/api/subjects")
def subjects() -> list[str]:
    return list_subjects()


@app.post("/api/subjects", status_code=201)
def create_subject_endpoint(request: CreateSubjectRequest) -> dict:
    subject_name = validate_new_subject_name(request.name)
    subject_root = (SUBJECTS_DIR.resolve() / subject_name).resolve()
    if subject_name in list_subjects() or subject_root.exists():
        raise HTTPException(status_code=409, detail=f"科目已存在：{subject_name}")

    try:
        subject_paths = create_subject(subject_name)
    except FileExistsError as exc:
        raise HTTPException(status_code=409, detail=f"科目已存在：{subject_name}") from exc
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    return {
        "name": subject_paths.name,
        "materials_dir": str(subject_paths.materials_dir),
        "outputs_dir": str(subject_paths.outputs_dir),
    }


@app.get("/api/subjects/{subject}/materials")
def subject_materials(subject: str) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        materials = list_materials(subject_paths)
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    indexed_paths, warning = indexed_material_paths(subject_paths)

    response = {
        "materials": [
            {
                "file_name": material.get("file_name"),
                "relative_path": material.get("relative_path"),
                "file_type": material.get("file_type"),
                "size_bytes": material.get("size_bytes"),
                "chapter": material.get("chapter"),
                "build_status": material_build_status(material, indexed_paths, warning),
                "conversion_status": material.get("conversion_status"),
                "converted_pptx": material.get("converted_pptx"),
            }
            for material in materials
        ],
    }
    if warning:
        response["warning"] = warning
    return response


@app.get("/api/subjects/{subject}/materials/page-image")
def material_page_image(subject: str, source_path: str = "", page_number: int | None = None) -> Response:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        material_path = resolve_preview_material_file(subject_paths, source_path)
    except HTTPException as exc:
        if exc.detail == "Invalid source_path.":
            raise HTTPException(status_code=400, detail="source_path 无效，请重新选择来源。") from exc
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    if not material_path.exists() or not material_path.is_file():
        raise HTTPException(status_code=404, detail="找不到该来源文件，可能已被移动、删除或重命名。")
    if page_number is None:
        raise HTTPException(status_code=400, detail="缺少页面编号或幻灯片编号。")
    if page_number < 1:
        raise HTTPException(status_code=400, detail="页面编号或幻灯片编号必须从 1 开始。")

    suffix = material_path.suffix.lower()
    if suffix not in {".pdf", ".ppt", ".pptx"}:
        raise HTTPException(status_code=400, detail="该文件类型暂不支持页面预览，仅支持 PDF/PPT/PPTX。")

    cache_dir = preview_cache_dir(subject_paths)
    image_cache_path = cache_dir / f"{preview_cache_key(source_path, material_path, page_number)}.png"
    if image_cache_path.exists():
        return Response(content=image_cache_path.read_bytes(), media_type="image/png")

    render_source = material_path
    page_label = "页码"
    keep_cache_paths = {image_cache_path}
    if suffix in {".ppt", ".pptx"}:
        page_label = "幻灯片编号"
        keep_cache_paths.add(presentation_pdf_cache_path(cache_dir, source_path, material_path))

    cleanup_old_preview_cache(cache_dir, keep_cache_paths)

    if suffix in {".ppt", ".pptx"}:
        render_source = convert_presentation_to_cached_pdf(subject_paths, source_path, material_path)

    image_bytes = render_pdf_page_to_png(render_source, page_number, page_label)
    image_cache_path.write_bytes(image_bytes)

    return Response(content=image_bytes, media_type="image/png")


@app.post("/api/subjects/{subject}/materials/upload")
async def upload_subject_materials(subject: str, files: list[UploadFile] = File(...)) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        upload_items = [
            ApiUploadedFile(file.filename or f"upload-{index}", await file.read())
            for index, file in enumerate(files, start=1)
        ]
        result = save_uploaded_materials(subject_paths, upload_items)
    except HTTPException:
        raise
    except ValueError as exc:
        return {
            "success": False,
            "saved_files": [],
            "message": "",
            "warning": str(exc),
            "error": str(exc),
        }
    except Exception as exc:
        return {
            "success": False,
            "saved_files": [],
            "message": "",
            "warning": f"上传失败：{exc}",
            "error": str(exc),
        }

    return {
        "success": True,
        "saved_files": result.get("saved") or [],
        "message": "上传成功，请添加/更新知识库。",
        "warning": None,
        "error": None,
    }


@app.post("/api/subjects/{subject}/materials/delete")
def delete_subject_material(subject: str, request: DeleteMaterialRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        relative_path = request.relative_path.strip()
        if not relative_path:
            return {
                "success": False,
                "message": "",
                "warning": "relative_path 不能为空。",
                "error": "empty_relative_path",
                "deleted_relative_path": "",
            }

        result = soft_delete_subject_material(subject_paths, relative_path)
        deleted_relative_path = (result.get("result") or {}).get("deleted_relative_path", relative_path)
    except HTTPException:
        raise
    except FileNotFoundError as exc:
        return {
            "success": False,
            "message": "",
            "warning": str(exc),
            "error": "file_not_found",
            "deleted_relative_path": request.relative_path,
        }
    except ValueError as exc:
        return {
            "success": False,
            "message": "",
            "warning": str(exc),
            "error": "invalid_relative_path",
            "deleted_relative_path": request.relative_path,
        }
    except Exception as exc:
        return {
            "success": False,
            "message": "",
            "warning": f"移除资料失败：{exc}",
            "error": str(exc),
            "deleted_relative_path": request.relative_path,
        }

    return {
        "success": True,
        "message": "资料已移入回收目录，请重新建库。",
        "deleted_relative_path": deleted_relative_path,
    }


@app.post("/api/subjects/{subject}/materials/rename")
def rename_subject_material_endpoint(subject: str, request: RenameMaterialRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        old_relative_path = request.old_relative_path.strip()
        new_relative_path = request.new_relative_path.strip()
        if not old_relative_path or not new_relative_path:
            return {
                "success": False,
                "message": "",
                "warning": "old_relative_path 和 new_relative_path 不能为空。",
                "error": "empty_relative_path",
                "material": None,
            }

        result = rename_subject_material(subject_paths, old_relative_path, new_relative_path)
        material = result.get("result") or {}
    except HTTPException:
        raise
    except FileNotFoundError as exc:
        return {
            "success": False,
            "message": "",
            "warning": str(exc),
            "error": "file_not_found",
            "material": None,
        }
    except FileExistsError as exc:
        return {
            "success": False,
            "message": "",
            "warning": str(exc),
            "error": "target_exists",
            "material": None,
        }
    except ValueError as exc:
        return {
            "success": False,
            "message": "",
            "warning": str(exc),
            "error": "invalid_relative_path",
            "material": None,
        }
    except Exception as exc:
        return {
            "success": False,
            "message": "",
            "warning": f"重命名资料失败：{exc}",
            "error": str(exc),
            "material": None,
        }

    return {
        "success": True,
        "message": "资料已重命名，请重新建库。",
        "material": material,
    }


@app.post("/api/subjects/{subject}/materials/convert-ppt")
def convert_subject_ppt_material_endpoint(subject: str, request: ConvertPptRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        relative_path = request.relative_path.strip()
        if not relative_path:
            return {
                "success": False,
                "message": "",
                "warning": "relative_path 不能为空。",
                "error": "empty_relative_path",
                "converted_pptx": "",
                "archived_original_ppt": "",
            }

        result = convert_subject_ppt_material(subject_paths, relative_path)
        conversion = result.get("result") or {}
    except HTTPException:
        raise
    except FileNotFoundError as exc:
        return {
            "success": False,
            "message": "",
            "warning": str(exc),
            "error": "file_not_found",
            "converted_pptx": "",
            "archived_original_ppt": "",
        }
    except ValueError as exc:
        return {
            "success": False,
            "message": "",
            "warning": str(exc),
            "error": "invalid_ppt_material",
            "converted_pptx": "",
            "archived_original_ppt": "",
        }
    except Exception as exc:
        return {
            "success": False,
            "message": "",
            "warning": f"PPT 转换失败：{exc}",
            "error": str(exc),
            "converted_pptx": "",
            "archived_original_ppt": "",
        }

    return {
        "success": True,
        "message": "已转换为 PPTX，原始 PPT 已归档。",
        "converted_pptx": conversion.get("converted_pptx") or "",
        "archived_original_ppt": conversion.get("archived_original_ppt") or "",
    }


def read_index_manifest(subject_paths) -> dict:
    manifest_path = subject_paths.outputs_dir / "index_manifest.json"
    if not manifest_path.exists():
        return {}
    try:
        with manifest_path.open("r", encoding="utf-8") as f:
            manifest = json.load(f)
            return {
                "manifest_file_count": manifest.get("file_count", 0),
                "manifest_chunk_count": manifest.get("chunk_count", 0),
            }
    except (json.JSONDecodeError, OSError):
        return {}


@app.get("/api/subjects/{subject}/status")
def subject_status(subject: str) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        stats = material_stats(subject_paths)
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    indexed_count, index_status, warning = index_health(subject_paths)
    manifest_info = read_index_manifest(subject_paths)

    response = {
        "subject": subject_paths.name,
        "file_count": stats.get("file_count", 0),
        "total_size_bytes": stats.get("total_size_bytes", 0),
        "indexed_count": indexed_count,
        "index_status": index_status,
        "deepseek_configured": is_deepseek_configured(),
        "materials_dir": str(subject_paths.materials_dir),
        "outputs_dir": str(subject_paths.outputs_dir),
    }
    if manifest_info:
        response["manifest_file_count"] = manifest_info["manifest_file_count"]
        response["manifest_chunk_count"] = manifest_info["manifest_chunk_count"]
        if (
            not warning
            and manifest_info["manifest_chunk_count"] > 0
            and indexed_count == 0
        ):
            response["chroma_mismatch"] = True
            response["warning"] = "已切块的文本数大于 0，但 Chroma 索引为空，请重新建库。"
    if warning:
        response["warning"] = warning
        if manifest_info and manifest_info["manifest_chunk_count"] > 0 and indexed_count == 0:
            response["chroma_mismatch"] = True
    return response


@app.post("/api/subjects/{subject}/qa")
def subject_qa(subject: str, request: QaRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        _, index_status, warning = index_health(subject_paths)
        if index_status == "corrupted":
            return qa_error_response(warning or corrupted_index_warning())
        indexed_sources = list_indexed_sources(outputs_dir=subject_paths.outputs_dir)
        indexed_paths = indexed_source_path_set(indexed_sources)
        result = ask_course_question(
            subject_paths,
            question=request.question,
            top_k=request.top_k,
            use_deepseek=request.use_deepseek,
            selected_sources=request.source_filters,
            indexed_paths=indexed_paths,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        return qa_error_response(str(exc))

    if not result.get("success"):
        return qa_error_response(str(result.get("error") or "查询未成功，请稍后重试。"))

    answer_text = result.get("answer") or ""
    hits = [api_hit(hit) for hit in result.get("hits", [])]

    try:
        create_qa_record(
            subject=subject,
            question=request.question,
            answer=answer_text,
            hits_count=len(hits),
            answer_mode=result.get("answer_mode") or "",
            source_filters=request.source_filters,
            warning=result.get("warning"),
            rewritten_query=result.get("rewritten_query") or "",
            hits=hits,
        )
    except Exception:
        pass

    return {
        "answer": answer_text,
        "warning": result.get("warning"),
        "hits": hits,
    }


@app.post("/api/subjects/{subject}/self-test")
def subject_self_test(subject: str, request: SelfTestRequest) -> dict:
    type_configs = [
        config.dict()
        for config in request.type_configs
        if config.count > 0
    ]
    total_count = sum(int(config["count"]) for config in type_configs)
    if not type_configs:
        raise HTTPException(status_code=400, detail="至少需要选择一种题型。")
    if total_count > 30:
        raise HTTPException(status_code=400, detail="自测题总题量最多 30 题。")

    try:
        subject_paths = resolve_existing_subject_paths(subject)
        _, index_status, warning = index_health(subject_paths)
        if index_status == "corrupted":
            return qa_error_response(warning or corrupted_index_warning())
        if index_status == "empty":
            return qa_error_response("当前科目还没有可用知识库，请先建立知识库。")

        result = generate_self_test(
            subject_paths,
            source_filters=request.source_filters or None,
            type_configs=type_configs,
            answer_mode=request.answer_mode,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("Self-test generation failed")
        return qa_error_response(str(exc))

    if not result.get("success"):
        return qa_error_response(str(result.get("error") or "自测题生成未成功，请稍后重试。"))

    answer_text = result.get("answer") or ""
    hits = [api_hit(hit) for hit in result.get("hits", [])]

    try:
        create_qa_record(
            subject=subject,
            question="生成自测题",
            answer=answer_text,
            hits_count=len(hits),
            answer_mode="自测题",
            source_filters=request.source_filters,
            warning=result.get("warning"),
            rewritten_query="",
            hits=hits,
        )
    except Exception:
        pass

    return {
        "answer": answer_text,
        "warning": result.get("warning"),
        "hits": hits,
    }


def snippet_keywords_response(request: SnippetKeywordsRequest) -> dict:
    text = request.text.strip()
    if not text:
        return {"keywords": [], "warning": None}
    if not load_ai_settings()["enabled"]:
        return {"keywords": [], "warning": AI_DISABLED_WARNING}

    try:
        return {"keywords": extract_snippet_keywords(text), "warning": None}
    except DeepSeekError as exc:
        return {
            "keywords": [],
            "warning": f"AI 关键词提取暂不可用，已使用本地规则高亮。{exc}",
        }
    except Exception as exc:
        return {
            "keywords": [],
            "warning": f"AI 关键词提取暂不可用，已使用本地规则高亮。{exc}",
        }


@app.get("/api/qa-history")
def qa_history_list(subject: str = "", limit: int = 50, offset: int = 0) -> dict:
    if not subject:
        return {"records": [], "total": 0}
    records = list_qa_records(subject, limit=limit, offset=offset)
    return {
        "records": [qa_history_to_dict(r) for r in records],
        "total": len(records),
    }


@app.get("/api/qa-history/{record_id}")
def qa_history_detail(record_id: int) -> dict:
    record = get_qa_record(record_id)
    if not record:
        raise HTTPException(status_code=404, detail="QA record not found.")
    return qa_history_to_dict(record)


@app.delete("/api/qa-history/{record_id}")
def qa_history_delete(record_id: int) -> dict:
    deleted = delete_qa_record(record_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="QA record not found.")
    return {"success": True, "deleted_id": record_id}


@app.post("/api/preview/snippet-keywords")
def preview_snippet_keywords(request: SnippetKeywordsRequest) -> dict:
    return snippet_keywords_response(request)


@app.post("/api/snippet-keywords")
def snippet_keywords(request: SnippetKeywordsRequest) -> dict:
    return snippet_keywords_response(request)


def _build_docx_response(
    title: str,
    subject: str,
    scope_label: str,
    generated_at: str,
    content: str,
    sources: list[dict],
    include_sources: bool,
    filename_prefix: str,
) -> Response:
    doc = Document()

    doc.add_heading(title, level=1)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run(f"科目：{subject}")
    run.font.size = Pt(10)

    if scope_label:
        meta2 = doc.add_paragraph()
        meta2.paragraph_format.space_after = Pt(2)
        run2 = meta2.add_run(f"资料范围：{scope_label}")
        run2.font.size = Pt(10)

    meta3 = doc.add_paragraph()
    meta3.paragraph_format.space_after = Pt(6)
    run3 = meta3.add_run(f"生成时间：{generated_at}")
    run3.font.size = Pt(10)

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(stripped)

    if include_sources and sources:
        doc.add_heading("参考来源", level=2)
        for index, source in enumerate(sources, start=1):
            metadata = source.get("metadata") or {}
            file_name = (
                metadata.get("file_name")
                or metadata.get("source_path")
                or "未知来源"
            )
            page = metadata.get("page_number")
            slide = metadata.get("slide_number")
            location = (
                f"第 {page} 页"
                if page
                else f"第 {slide} 张幻灯片"
                if slide
                else "位置未知"
            )

            ref = doc.add_paragraph()
            ref.paragraph_format.space_after = Pt(2)
            run_ref = ref.add_run(f"[{index}] {file_name}，{location}")
            run_ref.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = filename_prefix
    encoded_filename = quote(filename, safe='.')
    headers = {"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )



def _build_pdf_response(
    title: str,
    subject: str,
    scope_label: str,
    generated_at: str,
    content: str,
    sources: list[dict],
    include_sources: bool,
    filename_prefix: str,
) -> Response:
    # Build a DOCX in memory, convert it to PDF, and return a download.
    executable = find_libreoffice_executable()
    if not executable:
        raise HTTPException(
            status_code=500,
            detail="无法导出 PDF：未检测到 LibreOffice。",
        )

    requested_name = Path(filename_prefix or "export.pdf").name
    if not requested_name.lower().endswith(".pdf"):
        requested_name = f"{Path(requested_name).stem}.pdf"

    docx_response = _build_docx_response(
        title=title,
        subject=subject,
        scope_label=scope_label,
        generated_at=generated_at,
        content=content,
        sources=sources,
        include_sources=include_sources,
        filename_prefix="source.docx",
    )

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            input_docx = temp_root / "source.docx"
            output_dir = temp_root / "converted"
            output_dir.mkdir(parents=True, exist_ok=True)
            input_docx.write_bytes(bytes(docx_response.body))

            command = [
                executable,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(input_docx),
            ]
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=180,
            )

            if result.returncode != 0:
                logger.error(
                    "LibreOffice PDF export failed: stdout=%s stderr=%s",
                    result.stdout,
                    result.stderr,
                )
                raise HTTPException(
                    status_code=500,
                    detail="PDF 导出失败，请确认 LibreOffice 可正常运行。",
                )

            generated_pdf = output_dir / "source.pdf"
            if not generated_pdf.exists():
                candidates = list(output_dir.glob("*.pdf"))
                if not candidates:
                    raise HTTPException(
                        status_code=500,
                        detail="PDF 导出失败，未生成 PDF 文件。",
                    )
                generated_pdf = candidates[0]

            pdf_bytes = generated_pdf.read_bytes()

    except HTTPException:
        raise
    except subprocess.TimeoutExpired as exc:
        raise HTTPException(
            status_code=500,
            detail="PDF 导出超时，请稍后重试。",
        ) from exc
    except Exception as exc:
        logger.exception("PDF export failed")
        raise HTTPException(
            status_code=500,
            detail="PDF 导出失败，请检查后端日志。",
        ) from exc

    encoded_filename = quote(requested_name, safe=".")
    headers = {
        "Content-Disposition": (
            f"attachment; filename*=UTF-8''{encoded_filename}"
        )
    }

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers=headers,
    )


@app.post("/api/export/self-test/docx")
def export_self_test_docx(request: ExportSelfTestRequest) -> Response:
    return _build_docx_response(
        title=f"自测题 - {request.subject}",
        subject=request.subject,
        scope_label=request.scope_label,
        generated_at=request.generated_at,
        content=request.content,
        sources=request.sources,
        include_sources=request.include_sources,
        filename_prefix=request.filename or f"self-test-{request.subject}.docx",
    )


@app.post("/api/export/document/docx")
def export_document_docx(request: DocxExportRequest) -> Response:
    return _build_docx_response(
        title=request.title,
        subject=request.subject,
        scope_label=request.scope_label,
        generated_at=request.generated_at,
        content=request.content,
        sources=request.sources,
        include_sources=request.include_sources,
        filename_prefix=request.filename or request.filename_prefix,
    )



@app.post("/api/export/document/pdf")
def export_document_pdf(request: DocxExportRequest) -> Response:
    return _build_pdf_response(
        title=request.title,
        subject=request.subject,
        scope_label=request.scope_label,
        generated_at=request.generated_at,
        content=request.content,
        sources=request.sources,
        include_sources=request.include_sources,
        filename_prefix=(
            request.filename
            or request.filename_prefix
            or "export.pdf"
        ),
    )


@app.post("/api/subjects/{subject}/study-guide")
def subject_study_guide(subject: str, request: StudyGuideRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        _, index_status, warning = index_health(subject_paths)
        if index_status == "corrupted":
            return {
                "success": False,
                "content": "",
                "sources": [],
                "cached": False,
                "warning": warning or corrupted_index_warning(),
            }
        if index_status == "empty":
            return {
                "success": False,
                "content": "",
                "sources": [],
                "cached": False,
                "warning": "当前科目还没有可用知识库，请先建立知识库。",
            }

        result = generate_study_guide(
            subject_paths,
            unique_nonempty_files(request.source_filters),
            use_deepseek=request.use_deepseek,
            force_refresh=request.force_refresh,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        return {
            "success": False,
            "content": "",
            "sources": [],
            "cached": False,
            "warning": f"复习提纲生成失败：{exc}",
        }

    return {
        "success": bool(result.get("success")),
        "content": result.get("content") or "",
        "sources": result.get("sources") or [],
        "references": result.get("references") or [],
        "cached": bool(result.get("cached")),
        "warning": result.get("warning"),
    }


@app.post("/api/subjects/{subject}/overview")
def subject_overview(subject: str, request: OverviewRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        _, index_status, warning = index_health(subject_paths)
        if index_status == "corrupted":
            return {
                "success": False,
                "content": "",
                "sources": [],
                "cached": False,
                "warning": warning or corrupted_index_warning(),
            }
        if index_status == "empty":
            return {
                "success": False,
                "content": "",
                "sources": [],
                "cached": False,
                "warning": "当前科目还没有可用知识库，请先建立知识库。",
            }

        selected_sources = unique_nonempty_files(request.source_filters)
        if selected_sources:
            result = generate_current_material_overview(
                subject_paths,
                selected_sources,
                use_deepseek=request.use_deepseek,
                force_refresh=request.force_refresh,
            )
        else:
            result = generate_all_materials_overview(
                subject_paths,
                use_deepseek=request.use_deepseek,
                force_refresh=request.force_refresh,
            )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        return {
            "success": False,
            "content": "",
            "sources": [],
            "cached": False,
            "warning": f"资料概览生成失败：{exc}",
        }

    return {
        "success": bool(result.get("success")),
        "content": result.get("content") or "",
        "sources": result.get("sources") or [],
        "references": result.get("references") or [],
        "cached": bool(result.get("cached")),
        "warning": result.get("warning"),
    }


@app.post("/api/subjects/{subject}/longform")
def subject_longform(subject: str, request: LongformRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        _, index_status, warning = index_health(subject_paths)
        if index_status == "corrupted":
            return {
                "content": "",
                "outline": "",
                "group_summaries": [],
                "sources": [],
                "warnings": [warning or corrupted_index_warning()],
                "stats": {"total_chunks": 0, "used_chunks": 0, "groups_count": 0},
            }
        if index_status == "empty":
            return {
                "content": "",
                "outline": "",
                "group_summaries": [],
                "sources": [],
                "warnings": ["当前科目还没有可用知识库，请先建立知识库。"],
                "stats": {"total_chunks": 0, "used_chunks": 0, "groups_count": 0},
            }

        result = generate_longform_analysis(
            subject_paths,
            source_filters=request.source_filters or None,
            longform_type=request.longform_type,
            target_length=request.target_length,
            include_sources=request.include_sources,
            user_instruction=request.user_instruction,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("Longform 生成异常")
        return {
            "content": "",
            "outline": "",
            "group_summaries": [],
            "sources": [],
            "warnings": [f"Longform 生成失败：{exc}"],
            "stats": {"total_chunks": 0, "used_chunks": 0, "groups_count": 0},
        }

    content = str(result.get("content") or "").strip()

    if content:
        type_labels = {
            "analysis": "深度分析",
            "study_notes": "学习笔记",
            "report": "综合报告",
            "review": "读后感 / 心得体会",
            "outline": "提纲",
        }
        type_label = type_labels.get(
            request.longform_type,
            request.longform_type or "资料整理",
        )

        raw_sources = result.get("sources") or []
        history_sources = (
            [item for item in raw_sources if isinstance(item, dict)]
            if isinstance(raw_sources, list)
            else []
        )

        raw_warnings = result.get("warnings") or []
        history_warning = (
            "；".join(str(item) for item in raw_warnings if item)
            if isinstance(raw_warnings, list)
            else str(raw_warnings or "")
        )

        try:
            create_qa_record(
                subject=subject,
                question=f"资料整理｜{type_label}",
                answer=content,
                hits_count=len(history_sources),
                answer_mode="资料整理",
                source_filters=request.source_filters,
                warning=history_warning,
                rewritten_query=request.user_instruction or "",
                hits=history_sources,
            )
        except Exception:
            logger.exception(
                "Longform 已生成，但保存历史记录失败：subject=%s",
                subject,
            )

    return result


@app.post("/api/subjects/{subject}/index")
def subject_index(subject: str, request: IndexRequest) -> dict:
    try:
        subject_paths = resolve_existing_subject_paths(subject)
        if request.mode == "update":
            _, index_status, warning = index_health(subject_paths)
            if index_status == "corrupted":
                return index_error_response(
                    "当前索引已损坏，请使用重建当前范围知识库。",
                    error=warning,
                    error_type="index_corrupted",
                )

        selected_files: list[str] | None = None

        if request.scope == "selected":
            selected_files = unique_nonempty_files(request.files)
            if not selected_files:
                return index_error_response("请先选择要建库的资料。")

            available_paths = available_material_paths(subject_paths)
            missing_files = [file_name for file_name in selected_files if file_name not in available_paths]
            if missing_files:
                return index_error_response(
                    f"以下资料不存在：{'；'.join(missing_files)}。",
                    error="file_not_found",
                )
        else:
            selected_files = None

        if request.scope == "all" and request.mode == "reset":
            result = rebuild_all_materials_index(
                subject_paths,
                chunk_size=900,
                overlap=120,
                batch_size=32,
                embedding_model=None,
            )
        else:
            files_for_index = selected_files
            if request.scope == "all":
                files_for_index = sorted(available_material_paths(subject_paths))
                if not files_for_index:
                    return index_error_response("当前科目没有可建库的资料。")

            scope_label = "全部资料" if request.scope == "all" else f"{len(files_for_index or [])} 个资料"
            result = build_selected_materials_index(
                subject_paths,
                files_for_index or [],
                reset=request.mode == "reset",
                scope_label=scope_label,
                chunk_size=900,
                overlap=120,
                batch_size=32,
                embedding_model=None,
            )

        reset_chroma_client_cache(outputs_dir=subject_paths.outputs_dir)
        return index_success_response(result, request)
    except HTTPException:
        raise
    except Exception as exc:
        error = str(exc)
        logger.exception("建库异常: %s", error)
        if request.mode == "update" and is_index_corruption_error(error):
            return index_error_response(
                "当前索引已损坏，请使用重建当前范围知识库。",
                error=error,
                error_type="index_corrupted",
            )
        return index_error_response(
            "建库失败，请检查后端日志。",
            error=error,
        )
